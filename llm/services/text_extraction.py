"""
Serviço avançado de extração de texto com OCR fallback
"""

import asyncio
import io
import logging
from typing import Optional, Dict, Any
from pathlib import Path

from llm.exceptions import DocumentProcessingException

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Serviço avançado de extração de texto com OCR fallback"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_doc,
            '.txt': self._extract_from_txt,
            '.rtf': self._extract_from_rtf,
        }
        
    async def extract_text(
        self, 
        file_content: bytes, 
        filename: str,
        use_ocr_fallback: bool = True
    ) -> str:
        """
        Extrai texto de arquivo com fallback para OCR
        
        Args:
            file_content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            use_ocr_fallback: Se deve usar OCR como fallback
            
        Returns:
            Texto extraído do arquivo
            
        Raises:
            DocumentProcessingException: Se não conseguir extrair texto
        """
        
        logger.info(f"Extraindo texto de: {filename}")
        
        # Determinar extensão do arquivo
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise DocumentProcessingException(f"Formato não suportado: {file_extension}")
        
        try:
            # Tentar extração direta
            extractor = self.supported_formats[file_extension]
            text = await extractor(file_content)
            
            if text and text.strip():
                cleaned_text = self._clean_text(text)
                logger.info(f"Texto extraído com sucesso: {len(cleaned_text)} caracteres")
                return cleaned_text
            
            # Se não conseguiu extrair texto e OCR está habilitado
            if use_ocr_fallback and file_extension == '.pdf':
                logger.warning("Tentando extração com OCR...")
                return await self._extract_with_ocr(file_content)
                
        except Exception as e:
            logger.error(f"Erro na extração de texto: {str(e)}")
            
            # Último recurso: OCR para PDFs
            if use_ocr_fallback and file_extension == '.pdf':
                try:
                    return await self._extract_with_ocr(file_content)
                except Exception as ocr_error:
                    raise DocumentProcessingException(
                        f"Falha na extração de texto e OCR: {str(ocr_error)}"
                    )
            
            raise DocumentProcessingException(f"Erro na extração: {str(e)}")
        
        raise DocumentProcessingException("Não foi possível extrair texto do documento")
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extrai texto de PDF usando PyMuPDF"""
        try:
            import fitz  # PyMuPDF
            
            # Processar em thread separada para não bloquear
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(
                None, 
                self._extract_pdf_sync, 
                file_content
            )
            return text
            
        except ImportError:
            raise DocumentProcessingException("PyMuPDF não instalado")
    
    def _extract_pdf_sync(self, file_content: bytes) -> str:
        """Extração síncrona de PDF"""
        import fitz
        
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extrai texto de DOCX"""
        try:
            from docx import Document
            
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(
                None, 
                self._extract_docx_sync, 
                file_content
            )
            return text
            
        except ImportError:
            raise DocumentProcessingException("python-docx não instalado")
    
    def _extract_docx_sync(self, file_content: bytes) -> str:
        """Extração síncrona de DOCX"""
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Extrair texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    async def _extract_from_doc(self, file_content: bytes) -> str:
        """Extrai texto de DOC (formato antigo)"""
        # Para .doc, seria necessário uma lib específica como python-docx2txt
        # ou conversão via LibreOffice
        raise DocumentProcessingException(
            "Formato .doc não suportado diretamente. Converta para .docx"
        )
    
    async def _extract_from_txt(self, file_content: bytes) -> str:
        """Extrai texto de arquivo TXT"""
        try:
            # Tentar diferentes codificações
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingException("Não foi possível decodificar o arquivo")
            
        except Exception as e:
            raise DocumentProcessingException(f"Erro ao ler arquivo TXT: {str(e)}")
    
    async def _extract_from_rtf(self, file_content: bytes) -> str:
        """Extrai texto de RTF"""
        # RTF requer biblioteca específica
        raise DocumentProcessingException(
            "Formato RTF não suportado ainda"
        )
    
    async def _extract_with_ocr(self, file_content: bytes) -> str:
        """Extrai texto usando OCR (Tesseract)"""
        try:
            import fitz  # Para converter PDF em imagens
            from PIL import Image
            import pytesseract
            
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(
                None, 
                self._ocr_pdf_sync, 
                file_content
            )
            return text
            
        except ImportError as e:
            raise DocumentProcessingException(f"Dependências de OCR não instaladas: {str(e)}")
    
    def _ocr_pdf_sync(self, file_content: bytes) -> str:
        """OCR síncrono para PDF"""
        import fitz
        from PIL import Image
        import pytesseract
        
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        
        for page_num in range(min(len(doc), 10)):  # Limitar a 10 páginas por performance
            page = doc.load_page(page_num)
            
            # Converter página em imagem
            mat = fitz.Matrix(2.0, 2.0)  # Aumentar resolução
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Processar com OCR
            image = Image.open(io.BytesIO(img_data))
            text = pytesseract.image_to_string(image, lang='por')
            text_parts.append(text)
        
        doc.close()
        return "\n".join(text_parts)
    
    def _clean_text(self, text: str) -> str:
        """Limpeza e normalização do texto extraído"""
        
        if not text:
            return ""
        
        # Remove quebras de linha excessivas
        text = '\n'.join(line.strip() for line in text.split('\n'))
        
        # Remove linhas vazias múltiplas
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        # Remove espaços múltiplos
        while '  ' in text:
            text = text.replace('  ', ' ')
        
        return text.strip()
    
    async def get_document_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Obtém informações sobre o documento"""
        
        file_extension = Path(filename).suffix.lower()
        info = {
            "filename": filename,
            "extension": file_extension,
            "size_bytes": len(file_content),
            "supported": file_extension in self.supported_formats
        }
        
        # Informações específicas por tipo
        if file_extension == '.pdf':
            try:
                import fitz
                doc = fitz.open(stream=file_content, filetype="pdf")
                info.update({
                    "pages": len(doc),
                    "has_text": any(doc.load_page(i).get_text().strip() for i in range(min(3, len(doc)))),
                    "metadata": doc.metadata
                })
                doc.close()
            except:
                pass
        
        return info
